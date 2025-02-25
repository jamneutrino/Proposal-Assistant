from flask import Flask, render_template, request, jsonify, redirect, url_for, send_file, flash, abort, send_from_directory
from markupsafe import Markup
from flask_sqlalchemy import SQLAlchemy
from flask_migrate import Migrate
from datetime import datetime, timedelta
from sheets_helper import get_sheet_data
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import threading
import time
import io
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from dotenv import load_dotenv
import os
from werkzeug.utils import secure_filename
import json
import re
import requests
from flask_login import current_user, login_required
from flask_wtf.csrf import CSRFProtect, generate_csrf
import auth
from flask_limiter.errors import RateLimitExceeded
from flask_limiter.util import get_remote_address
import logging
from logging.handlers import RotatingFileHandler
import uuid
from flask_limiter import Limiter
from validation import validate_project_data, validate_item_data, validate_file_upload, ValidationError, validate_string

# Load environment variables from .env file
load_dotenv()

# Now you can access environment variables using os.getenv()
# Example:
# database_url = os.getenv('DATABASE_URL')
# google_client_id = os.getenv('GOOGLE_CLIENT_ID')

app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = os.getenv('DATABASE_URL', 'sqlite:///items.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SECRET_KEY'] = os.getenv('FLASK_SECRET_KEY', 'default-dev-key-please-change')
app.config['UPLOAD_FOLDER'] = 'static/item_images'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}

# CSP configuration
app.config['CSP_REPORT_ONLY'] = os.getenv('CSP_REPORT_ONLY', 'False').lower() == 'true'

# Initialize CSRF protection
csrf = CSRFProtect(app)

# Set up Content Security Policy
@app.after_request
def set_security_headers(response):
    """
    Add security headers to all responses to protect against XSS and other attacks.
    """
    # Content Security Policy
    csp = {
        'default-src': ["'self'"],
        'script-src': [
            "'self'", 
            "https://cdn.tailwindcss.com", 
            "https://cdn.jsdelivr.net", 
            "https://maps.googleapis.com",
            "'unsafe-inline'"
        ],
        'style-src': ["'self'", "https://fonts.googleapis.com", "'unsafe-inline'"],
        'font-src': ["'self'", "https://fonts.gstatic.com"],
        'img-src': ["'self'", "data:", "https://*.googleapis.com", "https://*.gstatic.com"],
        'connect-src': [
            "'self'", 
            "https://maps.googleapis.com",
            "https://www.googleapis.com"
        ],
        'frame-src': ["'none'"],
        'object-src': ["'none'"],
        'base-uri': ["'self'"],
        'form-action': ["'self'"],
        'report-to': ["'csp-endpoint'"],
        'report-uri': ["/csp-report"]
    }
    
    # Convert the CSP dictionary to a string
    csp_string = '; '.join([f"{key} {' '.join(value)}" for key, value in csp.items()])
    
    # Set the CSP header
    if app.config.get('CSP_REPORT_ONLY', False):
        response.headers['Content-Security-Policy-Report-Only'] = csp_string
    else:
        response.headers['Content-Security-Policy'] = csp_string
    
    # Add other security headers
    response.headers['X-Content-Type-Options'] = 'nosniff'  # Prevent MIME type sniffing
    response.headers['X-Frame-Options'] = 'DENY'  # Prevent clickjacking
    response.headers['X-XSS-Protection'] = '1; mode=block'  # Enable XSS filtering
    response.headers['Referrer-Policy'] = 'strict-origin-when-cross-origin'  # Control referrer information
    
    return response

# Register global error handler for rate limiting
@app.errorhandler(RateLimitExceeded)
def handle_rate_limit_exceeded(e):
    retry_after = 60  # Default value
    
    try:
        # Debug information
        print(f"Global handler - Rate limit exceeded: {e}")
        print(f"Global handler - Exception type: {type(e)}")
        print(f"Global handler - Request path: {request.path}")
        print(f"Global handler - Client IP: {get_remote_address()}")
        
        # Try to get the actual retry_after value from the exception
        if hasattr(e, 'retry_after'):
            print(f"Global handler - Found retry_after attribute: {e.retry_after}")
            retry_after = int(e.retry_after)
        elif hasattr(e, 'description'):
            print(f"Global handler - Found description attribute: {e.description}")
            # Try to extract from the description (e.g., "5 per minute")
            description = str(e.description)
            if 'minute' in description:
                # Extract the number from the description
                try:
                    limit_value = int(''.join(filter(str.isdigit, description.split('per')[0])))
                    retry_after = 60  # 1 minute in seconds
                    print(f"Global handler - Extracted from minute description: {retry_after}")
                except (ValueError, IndexError):
                    pass
            elif 'hour' in description:
                retry_after = 300  # 5 minutes in seconds
                print(f"Global handler - Using hour-based retry_after: {retry_after}")
            elif 'day' in description:
                retry_after = 600  # 10 minutes in seconds
                print(f"Global handler - Using day-based retry_after: {retry_after}")
    except Exception as ex:
        print(f"Error processing rate limit exception: {ex}")
        # Continue with default retry_after value
    
    try:
        print(f"Global handler - Final retry_after value: {retry_after}")
        
        # Return an error page directly instead of redirecting to avoid loops
        return render_template('rate_limit_error.html', 
                               message='Too many login attempts. Please try again later.',
                               retry_after=retry_after), 429
    except Exception as template_ex:
        print(f"Error rendering template: {template_ex}")
        
        # If template rendering fails, return a simple HTML response
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>Rate Limit Exceeded</title>
            <style>
                body {{ font-family: Arial, sans-serif; text-align: center; margin-top: 50px; }}
                h1 {{ color: #e53e3e; }}
                p {{ margin: 20px 0; }}
                .countdown {{ font-weight: bold; }}
                .btn {{ display: inline-block; padding: 8px 16px; background-color: #3182ce; 
                       color: white; text-decoration: none; border-radius: 4px; }}
            </style>
        </head>
        <body>
            <h1>Rate Limit Exceeded</h1>
            <p>Too many login attempts. Please try again in <span class="countdown">{retry_after}</span> seconds.</p>
            <p><a href="/" class="btn">Return to Home</a></p>
            <script>
                let seconds = {retry_after};
                const countdownElement = document.querySelector('.countdown');
                const countdown = setInterval(function() {{
                    seconds--;
                    countdownElement.textContent = seconds;
                    if (seconds <= 0) {{
                        clearInterval(countdown);
                        window.location.href = "/auth/login";
                    }}
                }}, 1000);
            </script>
        </body>
        </html>
        """
        return html, 429

# Make CSRF token available in all templates
@app.context_processor
def inject_csrf_token():
    return dict(csrf_token=generate_csrf(), csrf_meta=lambda: Markup('<meta name="csrf-token" content="' + generate_csrf() + '">'))

# Ensure upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Initialize SQLAlchemy
from models import db, User, Project, Item
db.init_app(app)
migrate = Migrate(app, db)

# Import and initialize auth
from auth import init_app
init_app(app)

# Get the limiter instance from auth
from auth import limiter

# Note: Manual rate limiting implementation has been removed.
# We're now using Flask-Limiter's decorator method in auth.py which is more reliable.

# Create a lock for thread-safe operations
cleanup_lock = threading.Lock()

# Function to clean up old temporary files
def cleanup_old_files():
    """Clean up temporary document files older than 15 minutes"""
    try:
        with cleanup_lock:
            docs_dir = 'static/generated_docs'
            if not os.path.exists(docs_dir):
                return
                
            current_time = time.time()
            fifteen_minutes_ago = current_time - 900  # 15 minutes in seconds
            
            # Get all files in the directory
            for filename in os.listdir(docs_dir):
                # Skip template files
                if filename in ['template.docx', 'reference.docx']:
                    continue
                    
                # Check if it's an output file
                if filename.startswith('output_'):
                    file_path = os.path.join(docs_dir, filename)
                    
                    # Check if it's a file and not a directory
                    if os.path.isfile(file_path):
                        # Get the file's modification time
                        file_mod_time = os.path.getmtime(file_path)
                        
                        # If the file is older than 15 minutes, delete it
                        if file_mod_time < fifteen_minutes_ago:
                            try:
                                os.remove(file_path)
                                print(f"Cleaned up old temporary file: {file_path}")
                            except Exception as e:
                                print(f"Error removing old file {file_path}: {str(e)}")
                                # Try again with a delay
                                try:
                                    time.sleep(0.5)
                                    if os.path.exists(file_path):
                                        os.remove(file_path)
                                        print(f"Cleaned up old temporary file on second attempt: {file_path}")
                                except Exception as e2:
                                    print(f"Failed to remove file on second attempt: {str(e2)}")
    except Exception as e:
        print(f"Error in cleanup_old_files: {str(e)}")

# Schedule periodic cleanup
def start_cleanup_scheduler():
    """Start a background thread to periodically clean up old files"""
    def run_cleanup():
        while True:
            cleanup_old_files()
            # Sleep for 15 minutes before next cleanup
            time.sleep(900)
    
    # Start the cleanup thread
    cleanup_thread = threading.Thread(target=run_cleanup, daemon=True)
    cleanup_thread.start()

# Start the cleanup scheduler when the app starts
start_cleanup_scheduler()

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# List of available items with their images
class ItemDefinition:
    def __init__(self, name, image_path=None):
        self.name = name
        self.image_path = image_path

ITEMS = []

def save_items_to_file():
    """Save the ITEMS list to a file"""
    with open('items.txt', 'w') as f:
        for item in ITEMS:
            f.write(f"{item.name}|{item.image_path or ''}\n")

def load_items_from_file():
    """Load items from file if it exists"""
    global ITEMS
    try:
        with open('items.txt', 'r') as f:
            ITEMS = []
            for line in f:
                if line.strip():
                    parts = line.strip().split('|')
                    name = parts[0]
                    image_path = parts[1] if len(parts) > 1 and parts[1] else None
                    ITEMS.append(ItemDefinition(name, image_path))
    except FileNotFoundError:
        # If file doesn't exist, save current items
        ITEMS = [ItemDefinition(name) for name in ['Curbs', 'Pipes']]
        save_items_to_file()

# Load items when app starts
load_items_from_file()

# Cache for prices
price_cache = {}
last_update = 0
CACHE_DURATION = 300  # 5 minutes in seconds

def update_price_cache():
    global price_cache, last_update
    current_time = time.time()
    
    # Update cache if it's expired
    if current_time - last_update > CACHE_DURATION:
        new_prices = get_sheet_data()
        if new_prices:  # Only update if we got data
            price_cache = new_prices
            last_update = current_time

def format_date(date_str):
    """Format a date string to MM-DD-YYYY format"""
    if not date_str:
        return ""
        
    # Try different formats
    formats = ['%Y-%m-%d', '%m/%d/%Y', '%m-%d-%Y', '%m/%d/%y', '%m-%d-%y']
    
    for fmt in formats:
        try:
            date_obj = datetime.strptime(date_str, fmt)
            return date_obj.strftime('%m-%d-%Y')
        except ValueError:
            continue
    
    # If all formats fail, return the original string
    return date_str

# Add a filter to convert MM-DD-YYYY to YYYY-MM-DD for HTML date inputs
@app.template_filter('format_date_for_input')
def format_date_for_input(date_str):
    """Convert MM-DD-YYYY to YYYY-MM-DD for HTML date inputs"""
    if not date_str:
        return ""
    
    try:
        # Try to parse the date as MM-DD-YYYY
        date_obj = datetime.strptime(date_str, '%m-%d-%Y')
        return date_obj.strftime('%Y-%m-%d')
    except ValueError:
        # Try other formats
        formats = ['%Y-%m-%d', '%m/%d/%Y', '%m-%d-%Y', '%m/%d/%y', '%m-%d-%y']
        
        for fmt in formats:
            try:
                date_obj = datetime.strptime(date_str, fmt)
                return date_obj.strftime('%Y-%m-%d')
            except ValueError:
                continue
        
        # If all formats fail, return empty string
        return ""

@app.route('/')
@login_required
def index():
    projects = Project.query.filter_by(user_id=current_user.id).order_by(Project.created_at.desc()).all()
    google_places_api_key = os.getenv('GOOGLE_PLACES_API_KEY')
    return render_template('index.html', projects=projects, google_places_api_key=google_places_api_key)

@app.route('/project/<int:project_id>')
@login_required
def project(project_id):
    project = Project.query.get_or_404(project_id)
    
    # Check if the project belongs to the current user
    if project.user_id != current_user.id and not current_user.is_admin:
        flash('You do not have permission to view this project.', 'error')
        return redirect(url_for('index'))
    
    update_price_cache()  # Update prices when viewing a project
    
    # Format the date to MM-DD-YYYY
    if project.date:
        project.date = format_date(project.date)
    
    translation = translate_to_words(project.items)
    
    # Pass the items list with their names and image paths
    items_with_images = [{'name': item.name, 'image_path': item.image_path} for item in ITEMS]
    
    # Get flash messages
    flash_messages = []
    from flask import get_flashed_messages
    messages = get_flashed_messages(with_categories=True)
    if messages:
        for category, message in messages:
            flash_messages.append({'category': category, 'message': message})
    
    return render_template('project.html', 
                         project=project, 
                         items=items_with_images, 
                         translation=translation,
                         price_cache=price_cache,
                         flash_messages=flash_messages)

@app.route('/get_price/<item_name>')
def get_price(item_name):
    update_price_cache()
    price = price_cache.get(item_name, None)
    return jsonify({'price': price})

@app.route('/create_project', methods=['POST'])
@login_required
def create_project():
    try:
        # Debug logging
        app.logger.debug(f"Form data received: {request.form}")
        
        # Check if project_name is in the form data
        if 'project_name' not in request.form or not request.form['project_name'].strip():
            error_msg = "Project name is required"
            app.logger.error(error_msg)
            flash(error_msg, 'error')
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({
                    'success': False,
                    'error': error_msg
                }), 400
            else:
                return redirect(url_for('index'))
        
        # Validate all project data
        validated_data = validate_project_data(request.form)
        
        # Debug logging
        app.logger.debug(f"Validated data: {validated_data}")
        
        # Create project with validated data
        project = Project(
            name=validated_data['name'],
            date=validated_data['date'],
            attn=validated_data['attn'],
            contractor_name=validated_data['contractor_name'],
            contractor_email=validated_data['contractor_email'],
            job_contact=validated_data['job_contact'],
            job_contact_phone=validated_data['job_contact_phone'],
            address=validated_data['address'],
            user_id=current_user.id
        )
        
        db.session.add(project)
        db.session.commit()
        
        # Debug logging
        app.logger.debug(f"Project created with ID: {project.id}")
        
        flash('Project created successfully!', 'success')
        
        # Check if this is an AJAX request
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return jsonify({
                'success': True,
                'redirect': url_for('project', project_id=project.id)
            })
        else:
            return redirect(url_for('project', project_id=project.id))
    except ValidationError as e:
        flash(str(e), 'error')
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return jsonify({
                'success': False,
                'error': str(e)
            }), 400
        else:
            return redirect(url_for('index'))
    except Exception as e:
        app.logger.error(f"Error creating project: {str(e)}")
        flash('An error occurred while creating the project.', 'error')
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return jsonify({
                'success': False,
                'error': 'An error occurred while creating the project.'
            }), 500
        else:
            return redirect(url_for('index'))

@app.route('/update_project/<int:project_id>', methods=['POST'])
@login_required
def update_project(project_id):
    project = Project.query.get_or_404(project_id)
    
    # Check if the project belongs to the current user
    if project.user_id != current_user.id and not current_user.is_admin:
        flash('You do not have permission to update this project.', 'error')
        return redirect(url_for('index'))
    
    try:
        # Validate all project data
        validated_data = validate_project_data(request.form)
        
        # Update project with validated data
        project.name = validated_data['name']
        project.date = validated_data['date']
        project.attn = validated_data['attn']
        project.contractor_name = validated_data['contractor_name']
        project.contractor_email = validated_data['contractor_email']
        project.job_contact = validated_data['job_contact']
        project.job_contact_phone = validated_data['job_contact_phone']
        project.address = validated_data['address']
        
        db.session.commit()
        
        flash('Project updated successfully!', 'success')
        return redirect(url_for('project', project_id=project.id, success=True))
    except ValidationError as e:
        flash(str(e), 'error')
        return redirect(url_for('project', project_id=project.id))
    except Exception as e:
        app.logger.error(f"Error updating project: {str(e)}")
        flash('An error occurred while updating the project.', 'error')
        return redirect(url_for('project', project_id=project.id))

@app.route('/add_item/<int:project_id>', methods=['POST'])
@login_required
def add_item(project_id):
    project = Project.query.get_or_404(project_id)
    
    # Check if the project belongs to the current user
    if project.user_id != current_user.id and not current_user.is_admin:
        return jsonify({'success': False, 'error': 'Permission denied'})
    
    try:
        data = request.get_json()
        
        # Validate item data
        validated_data = validate_item_data(data)
        
        item = Item(
            name=validated_data['name'], 
            quantity=validated_data['quantity'], 
            price=validated_data['price'], 
            project_id=project.id
        )
        
        db.session.add(item)
        db.session.commit()
        
        translation = translate_to_words(project.items)
        
        return jsonify({
            'success': True,
            'items': [{'name': i.name, 'quantity': i.quantity, 'price': i.price} for i in project.items],
            'translation': translation
        })
    except ValidationError as e:
        return jsonify({'success': False, 'error': str(e)}), 400
    except Exception as e:
        app.logger.error(f"Error adding item: {str(e)}")
        return jsonify({'success': False, 'error': 'An error occurred while adding the item'}), 500

@app.route('/clear_items/<int:project_id>', methods=['POST'])
@login_required
def clear_items(project_id):
    project = Project.query.get_or_404(project_id)
    
    # Check if the project belongs to the current user
    if project.user_id != current_user.id and not current_user.is_admin:
        return jsonify({'success': False, 'error': 'Permission denied'})
    
    for item in project.items:
        db.session.delete(item)
    db.session.commit()
    
    return jsonify({
        'success': True,
        'items': [],
        'translation': ''
    })

@app.route('/delete_project/<int:project_id>', methods=['POST'])
@login_required
def delete_project(project_id):
    project = Project.query.get_or_404(project_id)
    
    # Check if the project belongs to the current user
    if project.user_id != current_user.id and not current_user.is_admin:
        flash('You do not have permission to delete this project.', 'error')
        return redirect(url_for('index'))
    
    db.session.delete(project)
    db.session.commit()
    
    flash('Project deleted successfully!', 'success')
    return redirect(url_for('index'))

@app.route('/generate_word/<int:project_id>', methods=['POST'])
def generate_word(project_id):
    try:
        project = Project.query.get_or_404(project_id)
        
        # Create output directory if it doesn't exist
        os.makedirs('static/generated_docs', exist_ok=True)
        
        # Load the template
        template_path = os.path.join('static/generated_docs', 'template.docx')
        doc = Document(template_path)

        # Calculate total price of all items
        total_price = sum(item.quantity * item.price for item in project.items)
        
        # Process address for new placeholders
        street_address = ""
        city_address = ""
        if project.address:
            # Split address by commas
            parts = [p.strip() for p in project.address.split(',')]
            if len(parts) >= 3:  # Make sure we have enough parts
                street_address = parts[0]  # Everything before first comma
                # Take everything between first and last comma (city, state, zip)
                city_address = ', '.join(parts[1:-1])
        
        # Define placeholder mappings with variations
        placeholders = {
            '{{Name}}': project.name or '',
            '{{Date}}': format_date(project.date) if project.date else '',
            '{{Attn}}': project.attn or '',
            '{{ContractorName}}': project.contractor_name or '',
            '{{ContractorEmail}}': project.contractor_email or '',
            '{{JobContact}}': project.job_contact or '',
            '{{JobContactPhone}}': project.job_contact_phone or '',
            '{{StreetAdd}}': street_address,
            '{{CityAdd}}': city_address,
            '{{TotalPrice}}': f"${total_price:,.2f}" if total_price else '',
            # Also add variations with spaces
            '{{ Name }}': project.name or '',
            '{{ Date }}': format_date(project.date) if project.date else '',
            '{{ Attn }}': project.attn or '',
            '{{ ContractorName }}': project.contractor_name or '',
            '{{ ContractorEmail }}': project.contractor_email or '',
            '{{ JobContact }}': project.job_contact or '',
            '{{ JobContactPhone }}': project.job_contact_phone or '',
            '{{ StreetAdd }}': street_address,
            '{{ CityAdd }}': city_address,
            '{{ TotalPrice }}': f"${total_price:,.2f}" if total_price else '',
        }

        def find_placeholder(text):
            """Find any placeholder in the text"""
            for placeholder in placeholders:
                if placeholder in text:
                    print(f"Found placeholder: {placeholder} in text: {text}")  # Debug logging
                    return placeholder
            return None

        def process_text(text, formatting_info=None):
            """Process text and maintain formatting"""
            if not text:
                return text
                
            result = text
            print(f"Processing text: {text}")  # Debug logging
            
            # Normalize the text for comparison
            normalized_text = ''.join(text.split()).lower()  # Remove all whitespace and convert to lowercase
            
            # Check for contractor fields in normalized text
            if 'contractorname' in normalized_text:
                print(f"Found ContractorName in: {text}")  # Debug logging
                value = project.contractor_name or ''
                # Remove any curly braces around the value
                if value.startswith('{{') and value.endswith('}}'):
                    value = value[2:-2]
                return value
            
            if 'contractoremail' in normalized_text:
                print(f"Found ContractorEmail in: {text}")  # Debug logging
                value = project.contractor_email or ''
                # Remove any curly braces around the value
                if value.startswith('{{') and value.endswith('}}'):
                    value = value[2:-2]
                return value
            
            # Handle other placeholders
            for placeholder, value in placeholders.items():
                if placeholder in text:
                    print(f"Replacing {placeholder} with {value}")  # Debug logging
                    # Remove any curly braces around the value
                    if value.startswith('{{') and value.endswith('}}'):
                        value = value[2:-2]
                    result = result.replace(placeholder, value)
            
            return result

        def process_paragraph(paragraph):
            # First, collect all runs and their text
            runs = paragraph.runs
            if not runs:
                return

            # Get the full paragraph text and print it for debugging
            full_text = paragraph.text
            print(f"Processing paragraph text: {full_text}")  # Debug logging
            
            # Normalize the text for comparison
            normalized_text = ''.join(full_text.split()).lower()  # Remove whitespace and convert to lowercase
            
            # Special handling for contractor fields
            if 'contractorname' in normalized_text or 'contractoremail' in normalized_text:
                # Process each run individually to maintain formatting
                for run in runs:
                    normalized_run = ''.join(run.text.split()).lower()
                    if 'contractorname' in normalized_run:
                        value = project.contractor_name or ''
                        # Remove any curly braces around the value
                        if value.startswith('{{') and value.endswith('}}'):
                            value = value[2:-2]
                        run.text = value
                    elif 'contractoremail' in normalized_run:
                        value = project.contractor_email or ''
                        # Remove any curly braces around the value
                        if value.startswith('{{') and value.endswith('}}'):
                            value = value[2:-2]
                        run.text = value
                return

            # Check if there are any other placeholders in the full text
            has_placeholders = any(p in full_text for p in placeholders)
            if not has_placeholders:
                return

            # Store formatting information for each character position
            formatting = []
            current_pos = 0
            
            for run in runs:
                print(f"Run text: {run.text}")  # Debug logging
                for _ in range(len(run.text)):
                    formatting.append({
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font_size': run.font.size,
                        'font_name': run.font.name,
                        'style': run.style
                    })
                current_pos += len(run.text)

            # Replace all placeholders in the full text
            new_text = process_text(full_text)
            print(f"Text after replacement: {new_text}")  # Debug logging

            # Clear the paragraph
            paragraph.clear()

            # Add the text back with original formatting
            if len(formatting) > 0:
                current_format = formatting[0]
                current_text = ""

                for i, char in enumerate(new_text):
                    # If we have formatting information for this position and it's different
                    # from current format, create a new run
                    if i < len(formatting) and (
                        formatting[i]['bold'] != current_format['bold'] or
                        formatting[i]['italic'] != current_format['italic'] or
                        formatting[i]['underline'] != current_format['underline'] or
                        formatting[i]['font_size'] != current_format['font_size'] or
                        formatting[i]['font_name'] != current_format['font_name']
                    ):
                        if current_text:
                            run = paragraph.add_run(current_text)
                            apply_format(run, current_format)
                            current_text = ""
                        current_format = formatting[i]

                    current_text += char

                # Add any remaining text
                if current_text:
                    run = paragraph.add_run(current_text)
                    apply_format(run, current_format)

        def apply_format(run, format_dict):
            """Apply stored formatting to a run"""
            run.bold = format_dict['bold']
            run.italic = format_dict['italic']
            run.underline = format_dict['underline']
            if format_dict['font_size']:
                run.font.size = format_dict['font_size']
            if format_dict['font_name']:
                run.font.name = format_dict['font_name']
            if format_dict['style']:
                run.style = format_dict['style']

        def process_table_cell(cell):
            """Process all paragraphs in a table cell"""
            for paragraph in cell.paragraphs:
                process_paragraph(paragraph)

        def process_shapes():
            """Process text in shapes and textboxes"""
            # Get all shapes from the document
            for shape in doc.inline_shapes:
                try:
                    # Try to access the textframe if it exists
                    if hasattr(shape, '_inline') and hasattr(shape._inline, 'graphic'):
                        # Access the shape's text frame
                        text_frame = shape._inline.graphic.graphicData.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}txbx')
                        if text_frame is not None:
                            # Process each paragraph in the text frame
                            for paragraph in text_frame.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
                                # Get all text elements in this paragraph
                                text_elements = paragraph.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                                if text_elements:
                                    # Combine all text elements to get the full text
                                    full_text = ''.join(elem.text or '' for elem in text_elements)
                                    print(f"Processing textbox full text: {full_text}")  # Debug logging
                                    
                                    # Check for contractor fields
                                    if 'ContractorName' in full_text:
                                        new_text = project.contractor_name or ''
                                        # Set the text in the first element and clear the others
                                        text_elements[0].text = new_text
                                        for elem in text_elements[1:]:
                                            elem.text = ''
                                    elif 'ContractorEmail' in full_text:
                                        new_text = project.contractor_email or ''
                                        # Set the text in the first element and clear the others
                                        text_elements[0].text = new_text
                                        for elem in text_elements[1:]:
                                            elem.text = ''
                                    else:
                                        # Handle other placeholders
                                        new_text = process_text(full_text)
                                        text_elements[0].text = new_text
                                        for elem in text_elements[1:]:
                                            elem.text = ''
                except Exception as e:
                    print(f"Error processing shape: {str(e)}")
                    continue

            # Process floating shapes (if any)
            body = doc._body._body
            for shape in body.findall('.//w:drawing', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                try:
                    # Try to find textboxes in the shape
                    textboxes = shape.findall('.//wps:txbx//w:p', {
                        'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
                        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                    })
                    
                    for textbox in textboxes:
                        # Get all text elements in this textbox
                        text_elements = textbox.findall('.//w:t', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                        if text_elements:
                            # Combine all text elements to get the full text
                            full_text = ''.join(elem.text or '' for elem in text_elements)
                            print(f"Processing textbox full text: {full_text}")  # Debug logging
                            
                            # Check for contractor fields
                            if 'ContractorName' in full_text:
                                new_text = project.contractor_name or ''
                                # Set the text in the first element and clear the others
                                text_elements[0].text = new_text
                                for elem in text_elements[1:]:
                                    elem.text = ''
                            elif 'ContractorEmail' in full_text:
                                new_text = project.contractor_email or ''
                                # Set the text in the first element and clear the others
                                text_elements[0].text = new_text
                                for elem in text_elements[1:]:
                                    elem.text = ''
                            else:
                                # Handle other placeholders
                                new_text = process_text(full_text)
                                text_elements[0].text = new_text
                                for elem in text_elements[1:]:
                                    elem.text = ''
                except Exception as e:
                    print(f"Error processing floating shape: {str(e)}")
                    continue

        # Process all paragraphs in the main document
        for paragraph in doc.paragraphs:
            process_paragraph(paragraph)
        
        # Process all tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    process_table_cell(cell)

        # Process shapes and textboxes
        process_shapes()
        
        # Create output filename with timestamp
        output_filename = f'output_{project.name}_{int(time.time())}.docx'
        output_path = os.path.join('static/generated_docs', output_filename)
        
        # Save the modified document
        doc.save(output_path)
        
        # Send the generated file
        response = send_file(
            output_path,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'{project.name}.docx'
        )
        
        # Clean up the generated file after sending
        @response.call_on_close
        def cleanup():
            try:
                # Wait a short time to ensure the file is fully sent
                time.sleep(0.5)
                if os.path.exists(output_path):
                    os.remove(output_path)
                    print(f"Successfully cleaned up {output_path}")
                else:
                    print(f"File {output_path} already removed or does not exist")
            except Exception as e:
                print(f"Error cleaning up {output_path}: {str(e)}")
                # Try again after a delay in case the file is still in use
                try:
                    time.sleep(1)
                    if os.path.exists(output_path):
                        os.remove(output_path)
                        print(f"Successfully cleaned up {output_path} on second attempt")
                except Exception as e2:
                    print(f"Failed to clean up {output_path} on second attempt: {str(e2)}")
        
        # Schedule a delayed cleanup task as a backup
        def delayed_cleanup():
            time.sleep(5)  # Wait 5 seconds after response is sent
            try:
                if os.path.exists(output_path):
                    os.remove(output_path)
                    print(f"Successfully cleaned up {output_path} with delayed cleanup")
            except Exception as e:
                print(f"Error in delayed cleanup for {output_path}: {str(e)}")
        
        # Start the delayed cleanup in a separate thread
        cleanup_thread = threading.Thread(target=delayed_cleanup)
        cleanup_thread.daemon = True
        cleanup_thread.start()
        
        return response
        
    except Exception as e:
        print(f"Error generating Word document: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/admin')
@login_required
def admin():
    # Check if user is admin
    if not current_user.is_admin:
        flash('You do not have permission to access the admin area.', 'error')
        return redirect(url_for('index'))
    
    return render_template('admin.html', items=ITEMS)

@app.route('/admin/items/add', methods=['POST'])
def admin_add_item():
    global ITEMS
    name = request.form.get('name')
    
    try:
        # Validate the name
        name = validate_string(name, "Item name", max_length=100)
        
        if name in [item.name for item in ITEMS]:
            return jsonify({'success': False, 'error': 'Item with this name already exists'}), 400
        
        # Validate the image file
        image = validate_file_upload(
            request.files.get('image'), 
            field_name="Image", 
            allowed_extensions=ALLOWED_EXTENSIONS, 
            max_size_mb=16
        )
        
        image_path = None
        if image:
            filename = secure_filename(image.filename)
            # Add timestamp to filename to prevent duplicates
            filename = f"{int(time.time())}_{filename}"
            image_path = os.path.join('/static/item_images', filename)  # Add /static/ prefix
            image.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
        
        ITEMS.append(ItemDefinition(name, image_path))
        ITEMS.sort(key=lambda x: x.name)  # Keep items sorted by name
        save_items_to_file()
        return jsonify({'success': True})
    except ValidationError as e:
        return jsonify({'success': False, 'error': str(e)}), 400
    except Exception as e:
        app.logger.error(f"Error adding item: {str(e)}")
        return jsonify({'success': False, 'error': 'An error occurred while adding the item'}), 500

@app.route('/admin/items/edit', methods=['POST'])
def admin_edit_item():
    global ITEMS
    try:
        old_name = validate_string(request.form.get('oldName'), "Old item name", max_length=100)
        new_name = validate_string(request.form.get('newName'), "New item name", max_length=100)
        keep_image = request.form.get('keepImage') == 'true'
        
        old_item = next((item for item in ITEMS if item.name == old_name), None)
        if not old_item:
            return jsonify({'success': False, 'error': 'Item not found'}), 404
            
        # Only check for name conflict if the name is actually changing
        if new_name != old_name and new_name in [item.name for item in ITEMS if item.name != old_name]:
            return jsonify({'success': False, 'error': 'Name already exists'}), 400

        # Validate the image file
        image = validate_file_upload(
            request.files.get('image'), 
            field_name="Image", 
            allowed_extensions=ALLOWED_EXTENSIONS, 
            max_size_mb=16,
            required=False
        )
        
        # Handle image update
        image_path = old_item.image_path if keep_image else None
        if image:
            # Delete old image if exists
            if old_item.image_path:
                old_image_path = os.path.join(app.config['UPLOAD_FOLDER'], os.path.basename(old_item.image_path))
                if os.path.exists(old_image_path):
                    os.remove(old_image_path)
            
            filename = secure_filename(image.filename)
            filename = f"{int(time.time())}_{filename}"
            image_path = os.path.join('/static/item_images', filename)  # Add /static/ prefix
            image.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
        elif not keep_image and old_item.image_path:
            # Delete old image if exists and not keeping it
            old_image_path = os.path.join(app.config['UPLOAD_FOLDER'], os.path.basename(old_item.image_path))
            if os.path.exists(old_image_path):
                os.remove(old_image_path)
        
        # Update the item
        old_item.name = new_name
        old_item.image_path = image_path
        
        # Save changes
        ITEMS.sort(key=lambda x: x.name)  # Keep items sorted by name
        save_items_to_file()
        
        return jsonify({'success': True})
    except ValidationError as e:
        return jsonify({'success': False, 'error': str(e)}), 400
    except Exception as e:
        app.logger.error(f"Error editing item: {str(e)}")
        return jsonify({'success': False, 'error': 'An error occurred while editing the item'}), 500

@app.route('/admin/items/delete', methods=['POST'])
def admin_delete_item():
    global ITEMS
    try:
        name = validate_string(request.form.get('name'), "Item name", max_length=100)
        
        item_to_delete = next((item for item in ITEMS if item.name == name), None)
        if not item_to_delete:
            return jsonify({'success': False, 'error': 'Item not found'}), 404
        
        # Delete image file if exists
        if item_to_delete.image_path:
            image_path = os.path.join(app.config['UPLOAD_FOLDER'], os.path.basename(item_to_delete.image_path))
            if os.path.exists(image_path):
                os.remove(image_path)
        
        # Remove item from list
        ITEMS = [item for item in ITEMS if item.name != name]
        save_items_to_file()
        
        return jsonify({'success': True})
    except ValidationError as e:
        return jsonify({'success': False, 'error': str(e)}), 400
    except Exception as e:
        app.logger.error(f"Error deleting item: {str(e)}")
        return jsonify({'success': False, 'error': 'An error occurred while deleting the item'}), 500

# Function to immediately clean up all temporary files
def cleanup_all_temp_files():
    """Clean up all temporary document files regardless of age"""
    try:
        with cleanup_lock:
            docs_dir = 'static/generated_docs'
            if not os.path.exists(docs_dir):
                return
            
            files_cleaned = 0
            
            # Get all files in the directory
            for filename in os.listdir(docs_dir):
                # Skip template files
                if filename in ['template.docx', 'reference.docx']:
                    continue
                    
                # Check if it's an output file
                if filename.startswith('output_'):
                    file_path = os.path.join(docs_dir, filename)
                    
                    # Check if it's a file and not a directory
                    if os.path.isfile(file_path):
                        try:
                            os.remove(file_path)
                            files_cleaned += 1
                            print(f"Cleaned up temporary file: {file_path}")
                        except Exception as e:
                            print(f"Error removing file {file_path}: {str(e)}")
                            # Try again with a delay
                            try:
                                time.sleep(0.5)
                                if os.path.exists(file_path):
                                    os.remove(file_path)
                                    files_cleaned += 1
                                    print(f"Cleaned up temporary file on second attempt: {file_path}")
                            except Exception as e2:
                                print(f"Failed to remove file on second attempt: {str(e2)}")
            
            return files_cleaned
    except Exception as e:
        print(f"Error in cleanup_all_temp_files: {str(e)}")
        return 0

@app.route('/admin/cleanup_temp_files', methods=['POST'])
def admin_cleanup_temp_files():
    """Admin route to clean up all temporary files"""
    files_cleaned = cleanup_all_temp_files()
    return jsonify({'success': True, 'files_cleaned': files_cleaned})

@app.route('/get_items_info', methods=['GET'])
def get_items_info():
    """Return a list of all items with their info (name, image_path)"""
    try:
        # Convert the ITEMS list to a JSON-friendly format
        items_data = []
        for item in ITEMS:
            items_data.append({
                'name': item.name,
                'image_path': item.image_path
            })
        
        return jsonify({
            'success': True,
            'items': items_data
        })
    except Exception as e:
        print(f"Error getting items info: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        })

@app.route('/api/places/autocomplete', methods=['GET'])
def places_autocomplete():
    try:
        input_text = validate_string(
            request.args.get('input', ''), 
            "Search input", 
            min_length=3, 
            required=True
        )
        
        api_key = os.getenv('GOOGLE_PLACES_API_KEY')
        if not api_key:
            app.logger.error("Google Places API key not found")
            return jsonify({"error": "API configuration error", "predictions": []}), 500
            
        # Sanitize input for URL
        input_text = requests.utils.quote(input_text)
        url = f"https://maps.googleapis.com/maps/api/place/autocomplete/json?input={input_text}&key={api_key}"
        
        response = requests.get(url)
        if response.status_code != 200:
            app.logger.error(f"Google Places API error: {response.status_code} - {response.text}")
            return jsonify({"error": "Failed to fetch places", "predictions": []}), response.status_code
            
        return jsonify(response.json())
    except ValidationError as e:
        return jsonify({"error": str(e), "predictions": []}), 400
    except Exception as e:
        app.logger.error(f"Error fetching places: {str(e)}")
        return jsonify({"error": "Failed to fetch places", "predictions": []}), 500

def translate_to_words(items):
    if not items:
        return "No items selected."
    
    item_counts = {}
    for item in items:
        if item.name in item_counts:
            item_counts[item.name] += item.quantity
        else:
            item_counts[item.name] = item.quantity
    
    parts = []
    
    for item, count in item_counts.items():
        if item == "Curbs":
            curb_text = "Curb" if count == 1 else "Curbs"
            parts.append(f"- Tie-In / Flash ({count}) {curb_text} with roofing material compatible to existing material.")
        elif item == "Pipes":
            pipe_text = "pipe / penetration" if count == 1 else "pipes / penetrations"
            parts.append(f"- Tie-In / Flash ({count}) {pipe_text} with roofing material compatible to existing material.")
        elif item == "Item 1":
            parts.append(f"- {count} flashing {item}.")
        elif item == "Item 2":
            parts.append(f"- {count} {item} panels.")
        else:
            parts.append(f"- {count} {item}.")
    
    # Join with double line breaks
    if len(parts) > 1:
        return "\n\n".join(parts)
    else:
        return parts[0]

# CSP violation reporting endpoint
@app.route('/csp-report', methods=['POST'])
def csp_report():
    """
    Endpoint for CSP violation reports.
    """
    if request.content_type == 'application/csp-report':
        report = request.get_json()
        app.logger.warning(f"CSP Violation: {json.dumps(report)}")
    return '', 204  # No content response

@app.route('/csp-violations')
@login_required
def csp_violations():
    """
    Display CSP violations for administrators.
    """
    if not current_user.is_admin:
        flash('You do not have permission to access this page.', 'error')
        return redirect(url_for('index'))
    
    # In a real application, you would retrieve violations from a database
    # For this example, we'll just show a sample violation
    sample_violation = {
        "csp-report": {
            "document-uri": "http://example.com/page.html",
            "referrer": "",
            "violated-directive": "script-src-elem",
            "effective-directive": "script-src-elem",
            "original-policy": "default-src 'self'; script-src 'self'",
            "disposition": "enforce",
            "blocked-uri": "http://example.com/js/script.js",
            "line-number": 42,
            "column-number": 8,
            "source-file": "http://example.com/page.html",
            "status-code": 0,
            "script-sample": ""
        }
    }
    
    return render_template('csp_violation.html', violation=sample_violation)

# Initialize authentication after database setup
with app.app_context():
    # Create all tables if they don't exist
    db.create_all()
    # No need to initialize auth again, it's already done above

if __name__ == "__main__":
    # Development server
    app.run(host='0.0.0.0', port=8080, debug=False) 